#!/usr/bin/env python
"""IVD 临床试验方案 - Markdown 转 Word（v4）
修复记录：
- v1: 基础转换
- v2: HTML 表格解析、** 标记清理
- v3: body_buffer 合并、colspan/rowspan、OMML 公式
- v4: 标题宋体、_Toc 全局清理、首行缩进、BOM 过滤
"""
import sys, re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
try:
    from latex2mathml.converter import convert as latex_to_mathml
    HAS_LATEX2MATHML = True
except ImportError:
    HAS_LATEX2MATHML = False


# OMML namespace
MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


def _m(tag, children=None, text=None):
    """创建 OMML 元素，children 自动展开嵌套列表。"""
    el = OxmlElement(tag)
    if text:
        t = OxmlElement('m:t')
        t.text = text
        r = OxmlElement('m:r')
        r.append(t)
        el.append(r)
    if children:
        for c in _flatten(children):
            if c is not None:
                el.append(c)
    return el


def _flatten(items):
    """展开嵌套列表。"""
    for item in items:
        if isinstance(item, list):
            yield from _flatten(item)
        else:
            yield item


def mathml_to_omml(mathml_str):
    """将 MathML 转为 OMML oxml 元素列表（简化版）。"""
    from lxml import etree
    root = etree.fromstring(mathml_str.encode())
    ns = {'m': 'http://www.w3.org/1998/Math/MathML'}

    def convert_node(node):
        tag = node.tag.split('}')[-1] if '}' in node.tag else node.tag
        if tag in ('mi', 'mn', 'mo', 'mtext'):
            return _m('m:r', text=node.text or '')
        elif tag == 'mrow':
            children = [convert_node(c) for c in node]
            # 简单容器：返回子节点列表（展开）
            result = []
            for c in children:
                if isinstance(c, list):
                    result.extend(c)
                else:
                    result.append(c)
            return result
        elif tag == 'mfrac':
            num_children = convert_node(node[0])
            den_children = convert_node(node[1])
            num = _m('m:num', children=num_children if isinstance(num_children, list) else [num_children])
            den = _m('m:den', children=den_children if isinstance(den_children, list) else [den_children])
            return _m('m:f', children=[num, den])
        elif tag == 'msqrt':
            e = _m('m:e', children=[convert_node(node[0])])
            return _m('m:rad', children=[e])
        elif tag == 'msub':
            base = _m('m:e', children=[convert_node(node[0])])
            sub = _m('m:sub', children=[convert_node(node[1])])
            return _m('m:sSub', children=[base, sub])
        elif tag == 'msup':
            base = _m('m:e', children=[convert_node(node[0])])
            sup = _m('m:sup', children=[convert_node(node[1])])
            return _m('m:sSup', children=[base, sup])
        elif tag in ('msubsup',):
            base = _m('m:e', children=[convert_node(node[0])])
            sub = _m('m:sub', children=[convert_node(node[1])])
            sup = _m('m:sup', children=[convert_node(node[2])])
            return [_m('m:sSub', children=[base, sub]), _m('m:sSup', children=[_m('m:e'), sup])]
        elif tag == 'mstyle':
            return convert_node(node[0])
        elif tag == 'munderover':
            return convert_node(node[0])
        return None

    children = []
    for child in root:
        result = convert_node(child)
        if isinstance(result, list):
            children.extend(result)
        elif result is not None:
            children.append(result)
    return children


def contains_latex(text):
    """判断段落是否包含需要渲染的 LaTeX 公式。"""
    patterns = [
        r'\\\\frac', r'\\\\sqrt', r'\\\\alpha', r'\\\\beta',
        r'_\{', r'_([a-zA-Z0-9])',
        r'\\\\times', r'\\\\approx', r'\\\\pm',
        r'\$[^$]+\$',
    ]
    return any(re.search(p, text) for p in patterns)


def add_omath_paragraph(doc, latex_text, is_block=False):
    """将 LaTeX 公式渲染为 Word OMML 方程式。"""
    if not HAS_LATEX2MATHML:
        # fallback: plain text
        p = doc.add_paragraph()
        text = latex_text
        text = re.sub(r'\$\$', '', text)
        text = re.sub(r'\$([^$]+)\$', lambda m: m.group(1), text)
        run = p.add_run(text)
        set_run_font(run, size=12)
        return p

    try:
        # 清理 LaTeX 包裹符号
        clean = latex_text.strip()
        if clean.startswith('$$') and clean.endswith('$$'):
            clean = clean[2:-2].strip()
        elif clean.startswith('$') and clean.endswith('$'):
            clean = clean[1:-1].strip()
        # 也处理 inline 中混合文本的情况：提取公式部分
        # 如果全段是公式
        mathml = latex_to_mathml(clean)
        omml_children = mathml_to_omml(mathml)

        p = doc.add_paragraph()
        omath = OxmlElement('m:oMath')
        for c in omml_children:
            if c is not None:
                omath.append(c)
        omathpara = OxmlElement('m:oMathPara')
        omathpara.append(omath)
        # 附加到段落
        p._element.append(omathpara)
        return p
    except Exception as e:
        import traceback
        print(f'[WARN] OMML conversion failed: {e}', file=sys.stderr)
        traceback.print_exc(file=sys.stderr)
        p = doc.add_paragraph()
        run = p.add_run(latex_text)
        set_run_font(run, size=12)
        return p


CN_FONT = "宋体"
CN_FONT_BOLD = "宋体"  # 标题也用宋体（加粗），不用黑体


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            tag = qn(f"w:{edge}")
            elem = tcBorders.find(tag)
            if elem is None:
                elem = OxmlElement(f"w:{edge}")
                tcBorders.append(elem)
            elem.set(qn("w:val"), val)
            elem.set(qn("w:sz"), "6" if val == "single" else "0")
            elem.set(qn("w:color"), "000000")


def set_run_font(run, font_name=CN_FONT, size=12, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def add_heading_styled(doc, text, level):
    # 清理 markdown 标记
    text = re.sub(r'^\*\*|\*\*$', '', text)
    if level == 0:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, font_name=CN_FONT_BOLD, size=16, bold=True)
        return p
    elif level == 1:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(text)
        set_run_font(run, font_name=CN_FONT_BOLD, size=14, bold=True)
        return p
    elif level == 2:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(text)
        set_run_font(run, font_name=CN_FONT_BOLD, size=13, bold=True)
        return p
    else:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(text)
        set_run_font(run, font_name=CN_FONT_BOLD, size=12, bold=True)
        return p


def markdown_to_inline(text):
    """解析 markdown 行内标记为富文本片段。"""
    text = re.sub(r'\$\$', '', text)
    text = re.sub(r'\$([^$]+)\$', lambda m: m.group(1), text)
    result = []
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            result.append((part[2:-2], True, False, False))
        else:
            sub_parts = re.split(r'(\*[^*]+\*)', part)
            for sp in sub_parts:
                if sp.startswith("*") and sp.endswith("*") and len(sp) > 2:
                    result.append((sp[1:-1], False, True, False))
                else:
                    code_parts = re.split(r'(`[^`]+`)', sp)
                    for cp in code_parts:
                        if cp.startswith("`") and cp.endswith("`"):
                            result.append((cp[1:-1], False, False, True))
                        else:
                            result.append((cp, False, False, False))
    return result


def add_rich_paragraph(doc, text, style="Normal"):
    p = doc.add_paragraph()
    # 参考文献 `[1]` `[1-5]` 等顶格，不缩进
    if re.match(r'^\[\d+[\-\d]*\]', text.strip()):
        p.paragraph_format.first_line_indent = Cm(0)
    parts = markdown_to_inline(text)
    for p_text, bold, italic, code in parts:
        if not p_text:
            continue
        run = p.add_run(p_text)
        if code:
            set_run_font(run, font_name="Consolas", size=11, bold=bold)
        else:
            set_run_font(run, size=12, bold=bold, italic=italic)
    return p


def add_table_from_md(doc, rows):
    """将二维数组渲染为 Word 三线表。"""
    if not rows or len(rows) < 1:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row_data in enumerate(rows):
        for ci in range(min(len(row_data), cols)):
            cell = table.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            val = row_data[ci][:60]
            is_header = ri == 0
            run = p.add_run(val)
            set_run_font(run, bold=is_header, size=10 if is_header else 11)
            if ri == 0:
                set_cell_border(cell, top="single", bottom="single")
            elif ri == len(rows) - 1:
                set_cell_border(cell, bottom="single")


def parse_html_table(html_text):
    """从 HTML <table> 提取行数据。先清理干扰标签。"""
    html_text = re.sub(r'<a\s+id="_Toc[^"]*"\s*>\s*</a>', '', html_text, flags=re.IGNORECASE)
    html_text = re.sub(r'<a\s+[^>]*>', '', html_text)
    html_text = re.sub(r'</a>', '', html_text)
    rows = []
    tr_pattern = re.compile(r'<tr[^>]*>(.*?)</tr>', re.DOTALL | re.IGNORECASE)
    td_pattern = re.compile(r'<(?:td|th)[^>]*>', re.DOTALL | re.IGNORECASE)
    for tr_match in tr_pattern.finditer(html_text):
        row = []
        # 找到 td 或 th 的开始标签及其内容
        pos = 0
        tr_content = tr_match.group(1)
        for td_match in td_pattern.finditer(tr_content):
            tag = td_match.group(0).split()[0].strip('<>')
            # 检查 colspan/rowspan
            td_tag = td_match.group(0)
            colspan = 1
            cs = re.search(r'colspan\s*=\s*["\']?(\d+)["\']?', td_tag, re.IGNORECASE)
            if cs: colspan = int(cs.group(1))
            rowspan = 1
            rs = re.search(r'rowspan\s*=\s*["\']?(\d+)["\']?', td_tag, re.IGNORECASE)
            if rs: rowspan = int(rs.group(1))
            # 提取结束标签
            end_tag = td_tag[:2] + '/' + td_tag[1:].split('>')[0] + '>' if '/' in tag else f'</{tag}>'
            # 如果带有属性，构建通用 end tag
            end_pattern = re.compile(r'</' + re.escape(tag.split()[0]) + r'\s*>', re.IGNORECASE)
            end_m = end_pattern.search(tr_content, td_match.end())
            if end_m:
                cell_raw = tr_content[td_match.end():end_m.start()]
            else:
                cell_raw = ''
            cell_text = re.sub(r'<[^>]+>', '', cell_raw)
            cell_text = re.sub(r'\s+', ' ', cell_text).strip()
            for _ in range(colspan):
                row.append(cell_text)
            # rowspan 需要跨多行——延迟处理，简单起见当前行只占一列
        if row:
            rows.append(row)
    # 补齐列数不等的行
    if rows:
        max_cols = max(len(r) for r in rows)
        for r in rows:
            while len(r) < max_cols:
                r.append('')
    return rows


def convert(md_path, docx_path):
    Path(md_path).parent.mkdir(parents=True, exist_ok=True)
    md_str = Path(md_path).read_text(encoding="utf-8")
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = CN_FONT
    style.font.size = Pt(12)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), CN_FONT)
    rFonts.set(qn("w:hAnsi"), CN_FONT)
    rFonts.set(qn("w:eastAsia"), CN_FONT)

    # 正文首行缩进两个汉字字符（12pt * 2 = 24pt ≈ 0.74cm）
    style.paragraph_format.first_line_indent = Cm(0.74)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # 全局清理 _Toc 锚点标签（来自模板的 Word 目录锚点）
    md_str = re.sub(r'<a\s+id="_Toc\d+"\s*>\s*</a>', '', md_str)
    # 清理 BOM 字符
    md_str = md_str.replace('\ufeff', '')

    lines = md_str.split("\n")
    warnings = []
    prev_was_blank = True  # track consecutive blank lines (only need one break)

    def get_lvl(l):
        if l.startswith("# ") and not l.startswith("## "): return 0
        if l.startswith("## "): return 1
        if l.startswith("### "): return 2
        if l.startswith("#### "): return 3
        return -1

    # Pre-group: merge consecutive body lines into paragraphs
    # Use blank lines as paragraph boundaries
    i = 0
    body_buffer = []  # accumulate body text lines

    def flush_body():
        """写入累积的 body 段落（合并为一段）。"""
        nonlocal body_buffer
        if not body_buffer:
            return
        # 合并为一段（行间用空格连接）
        para_text = " ".join(line.strip() for line in body_buffer if line.strip())
        para_text = re.sub(r'\s+', ' ', para_text).strip()
        if para_text:
            # 如果包含 LaTeX 公式（含反斜杠+frac/sqrt 等），渲染为 OMML 方程式
            if chr(92) in para_text and any(cmd in para_text for cmd in ['frac', 'sqrt', 'alpha', 'beta', 'times', 'approx']):
                add_omath_paragraph(doc, para_text)
            else:
                add_rich_paragraph(doc, para_text)
        body_buffer = []

    while i < len(lines):
        line = lines[i]
        lvl = get_lvl(line)

        if "【" in line and "】" in line:
            warnings.append(f"L{i+1}: {line.strip()[:60]}")

        # HTML <table> — 在表格前后各留一个段落间距
        if "<table" in line.lower():
            flush_body()
            html_parts = []
            while i < len(lines) and "</table>" not in lines[i].lower():
                html_parts.append(lines[i])
                i += 1
            if i < len(lines):
                html_parts.append(lines[i])
            html = "\n".join(html_parts)
            rows = parse_html_table(html)
            if rows:
                add_table_from_md(doc, rows)
            i += 1
            prev_was_blank = True
            continue

        # 标题 — 前后各留一个段落间距
        if lvl >= 0:
            flush_body()
            text = re.sub(r"^#{1,6}\s+", "", line).strip()
            add_heading_styled(doc, text, lvl)
            prev_was_blank = True
            i += 1
            continue

        # 引用
        if line.startswith("> "):
            flush_body()
            p = doc.add_paragraph()
            run = p.add_run(line[2:].strip())
            set_run_font(run, italic=True, size=12)
            prev_was_blank = True
            i += 1
            continue

        # 无序列表
        if re.match(r"^\s*[-*]\s+", line):
            flush_body()
            t = re.sub(r"^\s*[-*]\s+", "", line)
            add_rich_paragraph(doc, t)
            prev_was_blank = True
            i += 1
            continue

        # 有序列表
        if re.match(r"^\s*\d+\.\s+", line):
            flush_body()
            t = re.sub(r"^\s*\d+\.\s+", "", line)
            add_rich_paragraph(doc, t)
            prev_was_blank = True
            i += 1
            continue

        # 水平线
        if re.match(r"^---+$", line.strip()):
            flush_body()
            p = doc.add_paragraph()
            run = p.add_run("─" * 50)
            set_run_font(run, size=10)
            prev_was_blank = True
            i += 1
            continue

        # 空行 = 段落边界
        if not line.strip():
            flush_body()
            prev_was_blank = True
            i += 1
            continue

        # 正文行：累积到 body_buffer，遇到段落边界再 flush
        body_buffer.append(line)
        i += 1

    # 收尾
    flush_body()

    doc.save(str(docx_path))
    if warnings:
        print(f"[WARN] {len(warnings)} placeholders:")
        for w in warnings[:5]:
            print(f"  {w}")
    print(f"[OK] Generated: {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python ivd_export_docx.py <input.md> <output.docx>")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])

#!/usr/bin/env python
"""将 .md 临床试验方案写入 .docx 模板（按格式誊抄）

用法: python ivd_fill_template.py <input.md> <template.docx> <output.docx>

保留模板全部样式（Heading 1/2/3 标题层级、Normal 正文格式、表格边框、页眉页脚），
按标题匹配将 .md 内容逐段填入模板对应位置，删除旧内容写入新内容。

表格处理：直接替换表格单元格内的文本（保留表格行数和格式）。
"""
import sys, re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def clean_md(text: str) -> str:
    """去除 markdown 标记符号。"""
    text = re.sub(r'\*\*', '', text)     # **bold** → bold
    text = re.sub(r'\*([^*]+)\*', r'\1', text)  # *italic* → italic
    text = re.sub(r'`([^`]+)`', r'\1', text)     # `code` → code
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # [text](url) → text
    return text.strip()


def split_md_into_sections(md_path: Path):
    """将 .md 按一级/二级标题拆分成结构化章节。
    
    返回: [(章节名, 级别, [段落文本]), ...]
    """
    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")
    sections = []
    current_name = "前言"
    current_level = 0
    current_paras = []

    for line in lines:
        if line.startswith("# "):
            if current_paras or current_name != "前言":
                sections.append((current_name, current_level, current_paras))
            current_name = line[2:].strip()
            current_level = 1
            current_paras = []
        elif line.startswith("## "):
            if current_paras or current_name != "前言":
                sections.append((current_name, current_level, current_paras))
            current_name = line[3:].strip()
            current_level = 2
            current_paras = []
        elif line.strip():
            current_paras.append(line.rstrip())
        else:
            # 空行作为分段符
            if current_paras and current_paras[-1] != "":
                current_paras.append("")

    if current_paras or current_name != "前言":
        sections.append((current_name, current_level, current_paras))

    return sections


def is_heading_para(text: str) -> bool:
    """判断模板中的段落是否是一个章节标题。"""
    # 匹配 "一、", "二、", "三、"..."(一)", "(二)" 等
    patterns = [
        r'^[一二三四五六七八九十]+、',               # 一、二、三、
        r'^（[一二三四五六七八九十]+）',             # （一）（二）
        r'^\(\d+\)',                               # (1) (2)
        r'^\d+\.',                                 # 1. 2.
        r'^[一二三四五六七八九十]+[、.]',
        r'^(说明|前言|背景|目的|方法|设计|统计|监查|数据|风险|质控|伦理|不良|偏离|保密|职责|参考|缩略|目录|附录)',
    ]
    for p in patterns:
        if re.search(p, text):
            return True
    return False


def find_section_boundaries(doc, start_idx: int = 0):
    """在模板中定位所有章节标题的位置。
    返回: [(段落索引, 标题文本), ...]
    """
    boundaries = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        # Heading 样式直接命中
        if "Heading" in style or "Title" in style:
            boundaries.append((i, text))
        elif is_heading_para(text) and len(text) < 80:
            boundaries.append((i, text))
    return boundaries


def match_section_to_template(md_name: str, template_paras):
    """判断 .md 节名是否匹配模板段落。"""
    md_name = clean_md(md_name)
    for idx, (pi, pt) in enumerate(template_paras):
        t = clean_md(pt)[:10]
        m = md_name[:10]
        # 提取数字部分做比较
        md_num = re.findall(r'[一二三四五六七八九十]+|\d+', md_name)
        tm_num = re.findall(r'[一二三四五六七八九十]+|\d+', clean_md(pt))
        if md_num and tm_num and md_num[0] == tm_num[0]:
            return idx
        # 关键词匹配
        keywords = md_name[:6]
        if keywords and keywords in pt:
            return idx
    return -1


def get_table_by_heading(doc, heading_idx: int):
    """获取 heading 后最近的一个 Word 表格。"""
    # 检查 heading 所在行是否在表格中
    para = doc.paragraphs[heading_idx]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if para in cell.paragraphs:
                    return table
    # 检查 heading 后的段落是否在表格中
    for j in range(heading_idx + 1, min(heading_idx + 10, len(doc.paragraphs))):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if doc.paragraphs[j] in cell.paragraphs:
                        return table
    return None


def fill_template(md_path, template_path, output_path, skip_cover=False):
    md_path = Path(md_path)
    template_path = Path(template_path)
    output_path = Path(output_path)

    # 1. 解析 .md
    md_sections = split_md_into_sections(md_path)
    print(f"[INFO] 解析 .md: {len(md_sections)} 个章节")

    # 2. 加载模板
    doc = Document(str(template_path))
    boundaries = find_section_boundaries(doc)
    print(f"[INFO] 模板章节: {len(boundaries)} 个标题")

    # 3. 构建 .md 节名到模板标题的映射
    matched = 0
    unmatched = []

    # 封面区域（段落 0-40）：替换固定占位
    if not skip_cover:
        for i, para in enumerate(doc.paragraphs[:50]):
            t = para.text
            for md_sec in md_sections:
                md_name = clean_md(md_sec[0])
                paras = md_sec[2]
                paras_text = "\n".join(paras)
                # 找申办者
                if "申办者信息" in md_name and "申办者：" in t:
                    # 从 .md 找申办者名称
                    for pl in paras:
                        if "深圳市新产业" in pl or "申办者" in pl:
                            para.clear()
                            run = para.add_run(
                                f"申办者：{clean_md(pl.replace('申办者', '').replace('：', '').strip())}"
                                if "申办者" not in pl else clean_md(pl)
                            )
                            # 保持现有样式
                            break
                # 产品名称
                elif "七项呼吸道病原体核酸检测试剂盒" in t and len(t) < 60:
                    for pl in paras:
                        if "产品全称" in pl:
                            val = pl.split("：")[-1].strip() if "：" in pl else pl
                            para.clear()
                            para.add_run(clean_md(val))
                            break
                # 方案编号
                elif "方案编号" in t and "186" in t:
                    for pl in paras:
                        if "方案编号" in pl:
                            val = pl.split("：")[-1].strip() if "：" in pl else "（留空）"
                            para.clear()
                            para.add_run(f"方案编号：{clean_md(val)}")
                            break
                # 版本日期
                elif "版本日期" in t and "2024" in t:
                    for pl in paras:
                        if "版本日期" in pl:
                            val = pl.split("：")[-1].strip() if "：" in pl else pl
                            para.clear()
                            para.add_run(f"方案版本号：V1.0 版本日期：{clean_md(val)}")
                            break

    # 4. 正文区域：按标题匹配替换
    for md_name, md_level, md_paras in md_sections:
        # 找到模板中对应的章节范围
        match_idx = match_section_to_template(md_name, boundaries)
        if match_idx < 0:
            unmatched.append(md_name)
            continue

        heading_pi, heading_text = boundaries[match_idx]
        # 确定替换范围：从 heading 的下一个段落，到下一个 heading
        next_pi = len(doc.paragraphs)
        if match_idx + 1 < len(boundaries):
            next_pi = boundaries[match_idx + 1][0]

        # 清空旧内容（从 heading 的下一个非空段落到下一个 heading 前）
        content_start = heading_pi + 1
        content_end = next_pi

        # 获取该范围内的表格，逐个填充
        tbl = get_table_by_heading(doc, heading_pi)
        if tbl and md_name in ["缩略语", "说明及历史修订"]:
            # 填充表格：从 .md 中提取表格式内容
            table_lines = [pl for pl in md_paras if pl.strip().startswith("|")]
            if table_lines:
                plain_lines = [clean_md(l) for l in table_lines if "|" in l]
                # 用表格数据替换模板表格
                data_rows = []
                for line in plain_lines:
                    row_data = [c.strip() for c in line.strip("|").split("|")]
                    if row_data:
                        data_rows.append(row_data)
                if data_rows and len(data_rows) > 1:
                    # data_rows[0] = 表头, data_rows[1:] = 数据
                    # 填充模板表格
                    max_rows = min(len(data_rows), len(tbl.rows))
                    for ri in range(max_rows):
                        for ci in range(min(len(data_rows[ri]), len(tbl.rows[ri].cells))):
                            cell = tbl.rows[ri].cells[ci]
                            c_para = cell.paragraphs[0]
                            c_para.clear()
                            c_para.add_run(clean_md(data_rows[ri][ci][:60]))

        # 收集 .md 正文（去标题行、去表格行）
        md_body = []
        for pl in md_paras:
            line = pl.strip()
            if not line or line.startswith("|") or line.startswith("!["):
                continue
            if line.startswith("#"):
                continue
            md_body.append(clean_md(line))

        if not md_body:
            matched += 1
            continue

        # 查找模板中 heading 后的第一个非空段落
        replace_start = -1
        for j in range(content_start, min(content_end, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip():
                replace_start = j
                break

        if replace_start < 0:
            matched += 1
            continue

        # 逐段填充分配：将 md_body 分配到可用的模板段落
        # 找到所有非空的模板段落
        tmpl_paras = []
        for j in range(replace_start, content_end):
            if j >= len(doc.paragraphs):
                break
            if doc.paragraphs[j].text.strip():
                tmpl_paras.append(j)

        for bi, body_line in enumerate(md_body):
            if bi < len(tmpl_paras):
                pi = tmpl_paras[bi]
                p = doc.paragraphs[pi]
                # 保留第一个 run 的格式
                if p.runs:
                    first_run = p.runs[0]
                    font_specs = {
                        'name': first_run.font.name,
                        'size': first_run.font.size.pt if first_run.font.size else 12,
                        'bold': first_run.font.bold,
                        'italic': first_run.font.italic,
                    }
                    # 清掉后续 runs
                    for run in p.runs[1:]:
                        run._element.getparent().remove(run._element)
                    first_run.text = body_line[:200]
                    # 恢复字体
                    if font_specs['name']:
                        first_run.font.name = font_specs['name']
                    if font_specs['size']:
                        first_run.font.size = Pt(font_specs['size'])
                    if font_specs['bold'] is not None:
                        first_run.font.bold = font_specs['bold']
                else:
                    p.clear()
                    p.add_run(body_line[:200])

        matched += 1

    # 保存
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"[OK] 已生成: {output_path}")
    print(f"[INFO] 已匹配: {matched}/{len(md_sections)} 章节")
    if unmatched:
        print(f"[WARN] 未匹配章节: {', '.join(unmatched[:5])}")


def main():
    if len(sys.argv) < 4:
        print("用法: python ivd_fill_template.py <input.md> <template.docx> <output.docx>")
        sys.exit(1)
    fill_template(sys.argv[1], sys.argv[2], sys.argv[3])


if __name__ == "__main__":
    main()
